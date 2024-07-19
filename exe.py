import streamlit as st
from docx import Document
from io import BytesIO
#from qa import chain
from docx.shared import Pt

#def get_answer(query):
#    response = chain.invoke(query)
#    return response['result']

# Function to set the margins of the document
def set_margins(section, left, right, top, bottom, header, footer):
    # Set the margins (in inches)
    section.left_margin = Pt(left * 72)   # 1 inch = 72 points
    section.right_margin = Pt(right * 72)
    section.top_margin = Pt(top * 72)
    section.bottom_margin = Pt(bottom * 72)
    section.header_distance = Pt(header * 72)
    section.footer_distance = Pt(footer * 72)

# Function to create a DOCX file from responses with Times New Roman and font size 11
def create_docx(headings, responses):
    doc = Document()
    
    # Set narrow margins and header/footer space for the document
    section = doc.sections[0]
    set_margins(section, left=0.75, right=0.75, top=1, bottom=1, header=0.5, footer=0.5)  # Example values in inches

    # Define a custom style for normal text
    styles = doc.styles
    normal_style = styles.add_style('CustomNormal', 1)  # 1 corresponds to paragraph style
    font = normal_style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)  # Optional: add space after paragraphs

    # Define a custom style for headings
    heading_style = styles.add_style('CustomHeading', 1)  # 1 corresponds to paragraph style
    font = heading_style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)  # Slightly larger size for headings
    heading_style.paragraph_format.space_after = Pt(12)  # Optional: add space after headings

    for heading, response in zip(headings, responses):
        # Add heading with custom style
        doc.add_heading(heading, level=1)
        heading_paragraph = doc.paragraphs[-1]  # Get the last added paragraph (the heading)
        heading_paragraph.style = 'CustomHeading'
        
        # Add response with custom style
        print(response)
        print(type(response))
        paragraphs = response['result'].split('\n')
        for paragraph in paragraphs:
            if paragraph.strip():  # Avoid empty paragraphs
                p = doc.add_paragraph(paragraph.strip(), style='CustomNormal')

    return doc

# Function to generate the document and handle Streamlit UI
def generate_document(get_answer):
    # List of questions and headings
    questions = [
        "1.A summary of 1.Business Overview- Include information such as the company's formation/incorporation date, headquarters location, business description, employee count, latest revenues, stock exchange listing and market capitalization, number of offices and locations, and details on their clients/customers",
        """2.A summary of Business Segment Overview
                Extract the revenue percentage of each component (verticals, products, segments, and sections) as a part of the total revenue.
                Performance: Evaluate the performance of each component by comparing the current year's sales/revenue and market share with the previous year's numbers.
                Sales Increase/Decrease explanation: Explain the causes of the increase or decrease in the performance of each component.
                precise numbers without wasting extra words.""",
        "3.A summary of Breakdown of sales and revenue by geography, specifying the percentage contribution of each region to the total sales.",
        "4.A summary of Summarize geographical data, such as workforce, clients, and offices, and outline the company's regional plans for expansion or reduction",
        "5.A summary of Analyze and explain regional sales fluctuations, including a geographical sales breakdown to identify sales trends",
        "6.A summary of Year-over-year sales increase or decline and reasons for the change.",
        "7.A summary of Summary of rationale & considerations (risks & mitigating factors)",
        "8.A summary of SWOT Analysis and be precise"
    ]

    headings = [
        "Business Overview",
        "Business Segment Overview",
        "Breakdown of Sales and Revenue by Geography",
        "Geographical Data and Regional Plans",
        "Analysis of Regional Sales Fluctuations",
        "Year-over-Year Sales Changes",
        "Rationale and Considerations",
        "SWOT Analysis"
    ]

    # Initialize progress bar
    progress = st.progress(0)
    num_questions = len(questions)
    
    # Generate responses
    responses = []
    for i, question in enumerate(questions):
        response = get_answer(question)
        responses.append(response)
        progress.progress((i + 1) / num_questions)  # Update progress

    # Create the DOCX file
    doc = create_docx(headings, responses)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Provide a preview of the DOCX file
    st.subheader("Preview of Generated Responses")
    # Optionally, use docx_preview(buffer, height=600) if preview functionality is implemented

    # Provide a download button for the DOCX file
    st.download_button(
        label="Download Responses",
        data=buffer,
        file_name="responses.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Finish progress bar
    progress.empty()

# Set the title of the app
#st.title("AI SUMMARIZER")

# Button to generate answers
#if st.button('Generate Responses'):
#    generate_document()